import argparse
from pathlib import Path
from docx import Document

def append_document(dst: Document, src: Document):
    # Append src body elements into dst
    for element in src.element.body:
        dst.element.body.append(element)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--main", required=True, help="Main dissertation .docx")
    ap.add_argument("--results", required=True, help="Results_Chapter.docx")
    ap.add_argument("--discussion", required=True, help="Discussion_Conclusion.docx")
    ap.add_argument("--out", required=True, help="Output merged .docx")
    args = ap.parse_args()

    main_path = Path(args.main).expanduser()
    results_path = Path(args.results).expanduser()
    discussion_path = Path(args.discussion).expanduser()
    out_path = Path(args.out).expanduser()

    if not main_path.exists():
        raise FileNotFoundError(f"Main doc not found: {main_path}")
    if not results_path.exists():
        raise FileNotFoundError(f"Results doc not found: {results_path}")
    if not discussion_path.exists():
        raise FileNotFoundError(f"Discussion doc not found: {discussion_path}")

    main_doc = Document(str(main_path))
    results_doc = Document(str(results_path))
    discussion_doc = Document(str(discussion_path))

    main_doc.add_page_break()
    append_document(main_doc, results_doc)

    main_doc.add_page_break()
    append_document(main_doc, discussion_doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    main_doc.save(str(out_path))
    print("[OK] Merged output:", out_path)

if __name__ == "__main__":
    main()
