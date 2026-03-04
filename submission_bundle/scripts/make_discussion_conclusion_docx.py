import os
from docx import Document
from docx.shared import Pt

BASE = os.path.expanduser("~/dissertation/final_outputs/results_pack")
OUT_DOCX = os.path.join(BASE, "Discussion_Conclusion.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

doc.add_heading("Discussion and Conclusion", level=1)

doc.add_heading("1. Discussion", level=2)
doc.add_paragraph(
    "The SUMO-based evaluation demonstrates the feasibility of integrating security processing, trust management, "
    "and blockchain-style logging into a V2X communication workflow. The normalized PDR remains stable across "
    "different densities and speeds under the tested conditions, indicating that the security mechanisms do not collapse "
    "the communication layer. Delay remains low because cryptographic operations are modeled as fixed processing delays."
)
doc.add_paragraph(
    "Trust and blockchain mechanisms introduce additional control-plane overhead (misbehavior reports and periodic commits). "
    "This overhead is justified because it enables detection and mitigation of misbehaving nodes, and supports trust-based "
    "handover decisions at RSUs. Threshold sweeps highlight a trade-off: higher fast thresholds reduce risk but may increase "
    "rejections; lower thresholds improve availability but reduce resilience."
)

doc.add_heading("2. Threats to Validity", level=2)
doc.add_paragraph(
    "Results depend on assumptions about wireless propagation, cryptographic delay modeling, and attacker behavior. "
    "Real deployments may observe different timing due to hardware and channel conditions. RSU placement and mobility routes "
    "also influence handover frequency."
)

doc.add_heading("3. Conclusion", level=2)
doc.add_paragraph(
    "This project presents a secure trust-based V2X communication model with blockchain support implemented in ns-3 "
    "and evaluated under SUMO mobility. The integrated approach shows that security checks, trust updates, and ledger commits "
    "can coexist with V2X data exchange while maintaining stable delivery performance."
)

doc.add_heading("4. Future Work", level=2)
doc.add_paragraph(
    "Future work includes: (i) integrating real cryptographic libraries or more detailed timing models; "
    "(ii) exploring more realistic RSU deployment strategies and multi-hop V2I relaying; "
    "(iii) adding additional attacker types (Sybil, message tampering, DoS); "
    "(iv) scaling to larger maps and longer simulation durations."
)

doc.save(OUT_DOCX)
print("[OK] Wrote:", OUT_DOCX)
