import os, glob
import pandas as pd
from docx import Document
from docx.shared import Pt

BASE = os.path.expanduser("~/dissertation/final_outputs/results_pack")
OUT_DOCX = os.path.join(BASE, "Results_Chapter.docx")
CAPTIONS_TXT = os.path.join(BASE, "Figure_Captions.txt")

SUMO_SUM = os.path.join(BASE, "sumo_pipeline_mean_std.csv")
SWEEP_SUM = os.path.join(BASE, "final_sweep_mean_std.csv")
PAPER_TABLE = os.path.join(BASE, "paper_table_default_thresholds.csv")

SUMO_PLOTS_DIR = os.path.join(BASE, "sumo_density_plots")
SWEEP_PLOTS_DIR = os.path.join(BASE, "sweep_plots")

def read_csv(path):
    return pd.read_csv(path) if os.path.exists(path) else None

sumo = read_csv(SUMO_SUM)
sweep = read_csv(SWEEP_SUM)
ptable = read_csv(PAPER_TABLE)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

doc.add_heading("Results and Evaluation", level=1)

doc.add_heading("1. Experimental Setup (SUMO + NS-3)", level=2)
doc.add_paragraph(
    "This chapter evaluates the proposed Secure Trust-Based V2X Communication with Blockchain Support using "
    "SUMO-generated mobility traces converted to NS-2 mobility for ns-3. For each condition, 5 random seeds "
    "were used and mean±std values were reported."
)
doc.add_paragraph(
    "Metrics include normalized packet delivery ratio (PDR), average end-to-end delay, throughput, replay/signature drops, "
    "ledger trust, blockchain commit latency, and RSU handover performance (count and delay)."
)

doc.add_heading("2. Density and Speed Impact (SUMO Mobility)", level=2)
if sumo is not None and not sumo.empty:
    dens = sorted(sumo["nVehicles"].unique().tolist())
    spds = sorted(sumo["speedTag"].unique().tolist())
    doc.add_paragraph(f"Vehicle densities tested: {dens}. Speeds tested: {spds} (m/s).")

    pdr_min = float(sumo["pdr_norm_mean"].min())
    pdr_max = float(sumo["pdr_norm_mean"].max())
    dly_min = float(sumo["avgDelay_s_mean"].min())
    dly_max = float(sumo["avgDelay_s_mean"].max())
    thr_min = float(sumo["throughput_bps_mean"].min())
    thr_max = float(sumo["throughput_bps_mean"].max())

    doc.add_paragraph(
        f"Across the tested grid, normalized PDR remained in the range {pdr_min:.3f}–{pdr_max:.3f}. "
        f"Average delay remained low (approximately {dly_min:.6f}–{dly_max:.6f} seconds). "
        f"Throughput scaled with network load ({thr_min:.0f}–{thr_max:.0f} bps)."
    )

    doc.add_paragraph("Table 1: Mean±Std summary (subset shown). Full table is saved in results_pack.")
    cols = [c for c in ["nVehicles","speedTag","pdr_norm_mean","pdr_norm_std","avgDelay_s_mean","avgDelay_s_std","throughput_bps_mean","throughput_bps_std"] if c in sumo.columns]
    view = sumo[cols].head(12)

    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for j,c in enumerate(cols):
        hdr[j].text = c
    for _,row in view.iterrows():
        cells = table.add_row().cells
        for j,c in enumerate(cols):
            cells[j].text = str(row[c])
else:
    doc.add_paragraph("SUMO summary CSV not found.")

doc.add_heading("3. Security Stress Tests (Malicious Rate + Threshold Trade-offs)", level=2)
if sweep is not None and not sweep.empty:
    doc.add_paragraph(
        "We sweep maliciousRate and trust thresholds (fast/min) to analyze robustness and access control behavior. "
        "This provides a trade-off surface between availability (FAST authentication) and security (rejecting low-trust nodes)."
    )
else:
    doc.add_paragraph("Final sweep summary CSV not found.")

doc.add_heading("4. Figure Captions", level=2)
captions = []

def add_figs(dirpath, prefix):
    if not os.path.isdir(dirpath):
        return
    figs = sorted(glob.glob(os.path.join(dirpath, "*.png")))
    for f in figs:
        name = os.path.basename(f)
        cap = f"{prefix}: {name}"
        captions.append((name, cap))

add_figs(SUMO_PLOTS_DIR, "SUMO density-speed")
add_figs(SWEEP_PLOTS_DIR, "Final sweep")

if captions:
    for i,(name, cap) in enumerate(captions, start=1):
        doc.add_paragraph(f"Figure {i}. {cap}")
else:
    doc.add_paragraph("No plot images found in results_pack plot directories.")

with open(CAPTIONS_TXT, "w") as f:
    for i,(name, cap) in enumerate(captions, start=1):
        f.write(f"Figure {i}. {cap} (file: {name})\n")

doc.save(OUT_DOCX)
print("[OK] Wrote:", OUT_DOCX)
print("[OK] Wrote:", CAPTIONS_TXT)
